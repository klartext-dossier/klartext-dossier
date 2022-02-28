import io
import logging

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Emu

from colour import Color


class WordWriter:

    def __init__(self, template=None):
        self.document = docx.Document(template)   
        self.bold = False
        self.italic = False
        self.underlined = False
        self.superscript = False
        self.subscript = False
        self.strike = False
        self.code = False
        self.character_style = None
        self.p = None
        self.list_styles = []
        self.color = None


    def parse_content(self, child, parent):
        if 'heading' == child.tag:
            self.parse_heading(child, parent)
        elif 'paragraph' == child.tag:
            self.parse_paragraph(child, parent)
        elif 'list' == child.tag:
            self.parse_list(child, parent)
        elif 'table' == child.tag:
            self.parse_table(child, parent)
        elif 'box' == child.tag:
            self.parse_box(child, parent)
        elif 'break' == child.tag:
            self.parse_break(child, parent)
        elif 'run' == child.tag:
            self.parse_run(child, parent)
        elif 'style' == child.tag:
            self.parse_style(child, parent)
        elif 'nav' == child.tag:
            self.parse_nav(child, parent)
        elif 'image' == child.tag:
            self.parse_image(child, parent)


    def parse_contents(self, tag, parent):
        for child in tag.iterchildren():
            self.parse_content(child, parent)


    def parse_image(self, tag, parent):
        if not tag.get('src').endswith('.png'):
            logging.warn(f'Ignoring {tag.get("src")} as the format is not known')
            return

        w = h = None
        width = tag.get('width')
        if width and width.endswith('%'):
            page_width = self.document.sections[0].page_width - self.document.sections[0].left_margin - self.document.sections[0].right_margin
            w = (int(width[:-1]) / 100.0) * page_width
        height = tag.get('height')
        if height and height.endswith('%'):
            page_height = self.document.sections[0].page_height - self.document.sections[0].top_margin - self.document.sections[0].bottom_margin
            h = (int(height[:-1]) / 100.0) * page_height

        parent.add_picture(tag.get('src'), width=w, height=h)


    def parse_nav(self, tag, parent):
        self.p = parent.add_paragraph()

        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = self.p.add_run()._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)


    def parse_heading(self, tag, parent):
        parent.add_heading(tag.text, level=int(tag.get('level')))


    def parse_style(self, tag, parent):
        bold = self.bold
        italic = self.italic
        underlined = self.underlined
        superscript = self.superscript
        subscript = self.subscript
        strike = self.strike
        code = self.code
        character_style = self.character_style
        color = self.color

        cl = tag.get('class')
        if 'b' == cl:
            self.bold = True
        elif 'i' == cl:
            self.italic = True
        elif 'ul' == cl:
            self.underlined = True
        elif 'sub' == cl:
            self.subscript = True
        elif 'sup' == cl:
            self.superscript = True
        elif 'strike' == cl:
            self.strike = True
        elif 'code' == cl:
            self.code = True
        else:
            self.character_style = cl

        if tag.get('color'):
            self.color = RGBColor.from_string(tag.get('color')[1:])

        self.parse_contents(tag, parent)

        self.bold = bold
        self.italic = italic
        self.underlined = underlined
        self.superscript = superscript
        self.subscript = subscript
        self.strike = strike
        self.code = code
        self.character_style = character_style
        self.color = color


    def parse_run(self, tag, parent):
        r = self.p.add_run(tag.text)
        r.bold = self.bold
        r.italic = self.italic
        r.underline = self.underlined
        r.font.superscript = self.superscript
        r.font.subscript = self.subscript
        r.font.strike = self.strike
        if self.code:
            r.font.name = "Courier New"
        if self.character_style:
            r.style = self.character_style
        if self.color:
            r.font.color.rgb = self.color
        if tag.get('style'):
            r.style = tag.get('style')


    def parse_paragraph(self, tag, parent):
        self.p = parent.add_paragraph(style=tag.get('style'))
        self.parse_contents(tag, parent)


    def parse_list(self, tag, parent):
        for child in tag.iterchildren():
            if 'item' == child.tag:
                tag.attrib['style'] = "List Bullet"
                self.parse_paragraph(tag, parent)                


    def parse_break(self, tag, parent):
        if 'page' == tag.get('class'):
            parent.add_page_break()
            

    def parse_table_paragraph(self, tag, parent, first_p):
        # there is a default paragraph in each cell, so do not add the first one
        if first_p:
            self.p = parent.paragraphs[0]
            if tag.get('style'):
                self.p.style = tag.get('style')
        else:
            self.p = parent.add_paragraph(style=tag.get('style'))
        self.parse_contents(tag, parent)


    def add_small_paragraph(self, parent):
        self.p = parent.add_paragraph()            
        self.p.paragraph_format.space_after = 0
        self.p.paragraph_format.space_before = 0

  
    def set_cell_bg_color(self, cell, color_hex):
        # https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
        tblCell = cell._tc
        tblCellProperties = tblCell.get_or_add_tcPr()
        clShading = OxmlElement('w:shd')
        clShading.set(qn('w:fill'), color_hex)
        tblCellProperties.append(clShading)


    def set_cell_margins(self, cell, **kwargs):
        # https://stackoverflow.com/questions/51060431/how-to-set-cell-margins-of-tables-in-ms-word-using-python-docx
        """
        cell:  actual cell instance you want to modify

        usage:

            set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

        provided values are in twentieths of a point (1/1440 of an inch).
        read more here: http://officeopenxml.com/WPtableCellMargins.php
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in [
            "top",
            "start",
            "bottom",
            "end",
        ]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)
        

    def set_cell_border(self, cell, **kwargs):
        # https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    def set_header_cell(self, cell):
        self.set_cell_bg_color(cell, "7DC1FF")


    def parse_box(self, tag, parent):
        table = parent.add_table(rows=1, cols=1)
        table.style = 'Normal Table'
        cell = table.cell(0,0)
        self.set_cell_margins(cell, top=50, bottom=50, left=50, right=50)
        if tag.get('color'):
            border = Color(tag.get('color'))
            bg = Color(tag.get('color'))
            bg.luminance = 0.8
            self.set_cell_bg_color(cell, bg.hex)
            self.set_cell_border(cell, start={"sz": 36, "val": "single", "color": border.hex, "space": "0"})
        first_p = True
        for child in tag.iterchildren():
            if 'run' == child.tag:
                self.parse_run(child, cell)
                first_p = False
            elif 'paragraph' == child.tag:
                self.parse_table_paragraph(child, cell, first_p)
                first_p = False
            else:
                self.parse_content(child, cell)
        self.add_small_paragraph(parent)


    def parse_table(self, tag, parent):
        rows = cols = 0
        for row in tag.iterchildren('tr'):
            rows += 1
            cols = max(cols, sum(int(col.get('colspan', '1')) for col in row.iterchildren()))
        
        table = parent.add_table(rows=rows, cols=cols)
        table.style = tag.get('style')


        r = 0
        for row in tag.iterchildren('tr'):
            c = 0
            for col in row.iterchildren():
                cell = table.cell(r,c)
                self.set_cell_margins(cell, top=50, bottom=50, left=50, right=50)
                self.p = cell.paragraphs[0]
                if 'th' == col.tag:
                    self.set_header_cell(cell)
                first_p = True
                for child in col.iterchildren():
                    if 'run' == child.tag:
                        self.parse_run(child, cell)
                        first_p = False
                    elif 'paragraph' == child.tag:
                        self.parse_table_paragraph(child, cell, first_p)
                        first_p = False
                    elif 'table' == child.tag:
                        self.parse_table(child, cell)
                    else:
                        self.parse_content(child, parent)

                colspan = int(col.get('colspan', '1'))
                if colspan > 1:
                    cell.merge(table.cell(r, c+colspan-1))
                c += colspan
            r += 1 

        # calculate table width
        full_width = sum(col.width for col in table.columns)
        logging.debug(f"Table width: {Emu(full_width).cm} cm")
        
        # determine the cell widths
        cell_widths = [None for i in range(len(table.columns))]
        for row in tag.iterchildren('tr'):
            c=0
            for col in row.iterchildren():
                if col.tag in ['th', 'td']:
                    if col.get('width'):
                        cell_widths[c] = col.get('width')
                c += 1
        logging.debug(f"Cell widths: {cell_widths}")

        # map to real lengths
        for i in range(len(table.columns)):
            if cell_widths[i] and cell_widths[i].endswith('%'):
                cell_widths[i] = float(cell_widths[i][:-1]) * full_width / 100.0
        logging.debug(f"Cell widths: {cell_widths}")

        # set the column widths
        for i in range(len(table.columns)):
            for cell in table.column_cells(i):
                if cell_widths[i]:
                    cell.width = cell_widths[i]

        self.add_small_paragraph(parent)


    def parse_document(self, tag):
        self.parse_contents(tag, self.document)


    def convert(self, xml):        
        for s in self.document.styles:
            logging.debug(f'word style {s.type}: {s.name}')
            
        root = xml.getroot()
        if 'document' == root.tag:
            self.parse_document(root)


        target_stream = io.BytesIO()
        self.document.save(target_stream)

        return target_stream.getbuffer()